"""
Cantio - Importer
Supports: TXT, DOCX, PDF, VideoPsalm (.json/.xml/.vpc), EasyWorship (.ewsx/.db), BibleShow (.bib)
"""
import os
import re
import json
import sqlite3
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path


# ── Encoding detection ────────────────────────────────────────────────────────

def detect_encoding(filepath):
    """Try a sequence of encodings and return the first one that reads successfully."""
    with open(filepath, "rb") as f:
        raw_start = f.read(4)

    # BOM sniffing
    if raw_start[:2] in (b"\xff\xfe", b"\xfe\xff"):
        return "utf-16"
    if raw_start[:3] == b"\xef\xbb\xbf":
        return "utf-8-sig"
    if raw_start[:4] in (b"\x00\x00\xfe\xff", b"\xff\xfe\x00\x00"):
        return "utf-32"

    for enc in ["utf-8", "utf-8-sig", "cp1250", "cp1252", "iso-8859-2", "latin-1"]:
        try:
            with open(filepath, "r", encoding=enc) as f:
                content = f.read(2000)
            if content and len(content) > 10:
                return enc
        except (UnicodeDecodeError, UnicodeError):
            continue
    return "latin-1"


def _try_parse_xml_with_encodings(filepath_or_bytes):
    """
    Parse XML from a filepath or raw bytes, trying multiple encodings.
    Returns the root Element or None.
    """
    if isinstance(filepath_or_bytes, (str, Path)):
        # Standard parse first — handles UTF-8, UTF-16 with BOM, and <?xml ...?> declarations
        try:
            return ET.parse(str(filepath_or_bytes)).getroot()
        except ET.ParseError:
            pass
        with open(filepath_or_bytes, "rb") as f:
            raw = f.read()
    else:
        raw = filepath_or_bytes
        try:
            return ET.fromstring(raw)
        except ET.ParseError:
            pass

    for enc in ["utf-16", "utf-16-le", "utf-16-be", "utf-8-sig",
                "utf-8", "cp1250", "cp1252", "iso-8859-2", "latin-1"]:
        try:
            text = raw.decode(enc)
            text = text.lstrip("\ufeff")  # strip BOM
            # Ensure there's an XML declaration
            if not text.strip().startswith("<?xml"):
                text = '<?xml version="1.0" encoding="utf-8"?>' + text
            root = ET.fromstring(text.encode("utf-8"))
            print(f"[XML] Parsed with encoding: {enc}")
            return root
        except Exception:
            continue
    return None


# ── Helpers ───────────────────────────────────────────────────────────────────

def text_to_slides(text):
    """Split text into slides by double newline (strophes)."""
    if not text or not str(text).strip():
        return []
    blocks = [b.strip() for b in str(text).strip().split("\n\n") if b.strip()]
    return blocks if blocks else [str(text).strip()]


def strip_rtf(text):
    """
    Extract plain text from an RTF-encoded string.
    Handles: metadata groups (fonttbl/colortbl/…), \\uN unicode, \\'xx hex, \\par breaks.
    """
    if not isinstance(text, str):
        text = str(text)
    text = text.strip()
    if not text.startswith("{\\rtf"):
        return text

    # 0. Remove known non-text metadata groups using a simple pass
    #    We iterate because groups may be nested one level deep
    _JUNK = ("fonttbl", "colortbl", "stylesheet", "info",
              "pict", "object", "filetbl", "listtable", "mmathPr")
    # Pattern that handles one level of nested braces inside the group
    _inner = r"[^{}]*(?:\{[^{}]*\}[^{}]*)*"
    for _ in range(3):  # a few passes handles shallow nesting
        for grp in _JUNK:
            text = re.sub(r"\{\\" + grp + _inner + r"\}", "", text)
            text = re.sub(r"\{\\[*]\s*\\" + grp + _inner + r"\}", "", text)

    # 1. RTF Unicode escapes: \uN? (N is signed decimal codepoint, ? is fallback char)
    def rtf_uni(m):
        try:
            n = int(m.group(1))
            if n < 0:
                n += 65536
            return chr(n)
        except Exception:
            return ""
    # Consume \uN plus its one-character ANSI fallback (may be \'xx or a plain char)
    text = re.sub(r"\\u(-?\d+)\s?(?:\\'[0-9a-fA-F]{2}|[^\\\s{}])?", rtf_uni, text)

    # 2. Hex-escaped characters (cp1252 default code page)
    def unhex(m):
        try:
            return bytes.fromhex(m.group(1)).decode("cp1252", errors="replace")
        except Exception:
            return ""
    text = re.sub(r"\\'([0-9a-fA-F]{2})", unhex, text)

    # 3. Paragraph / line breaks — BEFORE control-word removal
    text = re.sub(r"\\pard?\b[^a-zA-Z]?", "\n", text)
    text = re.sub(r"\\par\b\s?", "\n", text)
    text = re.sub(r"\\line\b\s?", "\n", text)
    text = re.sub(r"\\tab\b\s?", "\t", text)

    # 4. Remove remaining RTF control words (\word, \word123, \word-123)
    text = re.sub(r"\\[a-zA-Z]+\-?\d*\s?", "", text)

    # 5. Remove brace markers (structural, not content in RTF)
    text = re.sub(r"[{}\\]", "", text)

    # 6. Tidy whitespace
    text = re.sub(r" {2,}", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _has_garbage(text, threshold=0.15):
    """Return True if more than threshold fraction of chars look like encoding garbage."""
    if not text:
        return False
    garbage = sum(1 for c in text if ord(c) > 0xFFFD or c in "\x00\x01\x02\x03\x04\x05")
    return (garbage / len(text)) > threshold


# ── TXT ───────────────────────────────────────────────────────────────────────

def import_txt(filepath):
    enc = detect_encoding(filepath)
    with open(filepath, "r", encoding=enc, errors="replace") as f:
        content = f.read()
    title = Path(filepath).stem
    slides = text_to_slides(content)
    return [{"title": title, "content": content, "slides": slides, "language": "ro"}]


# ── DOCX ──────────────────────────────────────────────────────────────────────

def import_docx(filepath):
    try:
        import docx
        doc = docx.Document(filepath)
        content = "\n".join(p.text for p in doc.paragraphs)
        title = Path(filepath).stem
        slides = text_to_slides(content)
        return [{"title": title, "content": content, "slides": slides, "language": "ro"}]
    except Exception as e:
        raise ImportError(f"DOCX import failed: {e}")


# ── PDF ───────────────────────────────────────────────────────────────────────

def import_pdf(filepath):
    try:
        import fitz
        doc = fitz.open(filepath)
        content = "\n\n".join(page.get_text() for page in doc)
        doc.close()
        title = Path(filepath).stem
        slides = text_to_slides(content)
        return [{"title": title, "content": content.strip(), "slides": slides, "language": "ro"}]
    except Exception as e:
        raise ImportError(f"PDF import failed: {e}")


# ── VideoPsalm .json ──────────────────────────────────────────────────────────

def import_videopasalm_json(filepath):
    enc = detect_encoding(filepath)
    with open(filepath, "r", encoding=enc, errors="replace") as f:
        data = json.load(f)

    if isinstance(data, list):
        items = data
    elif isinstance(data, dict) and "songs" in data:
        items = data["songs"]
    elif isinstance(data, dict):
        items = [data]
    else:
        items = []

    songs = []
    for item in items:
        title = item.get("title", item.get("name", "Untitled"))
        language = item.get("language", item.get("lang", "ro"))
        verses = item.get("verses", item.get("strophes", item.get("parts", [])))
        slides = []
        for verse in verses:
            if isinstance(verse, dict):
                text = verse.get("text", verse.get("content", verse.get("lines", "")))
                if isinstance(text, list):
                    text = "\n".join(text)
            elif isinstance(verse, str):
                text = verse
            else:
                continue
            if text.strip():
                slides.append(text.strip())
        if not slides:
            slides = text_to_slides(item.get("content", item.get("text", "")))
        content = "\n\n".join(slides)
        songs.append({
            "title": title,
            "content": content,
            "slides": slides,
            "language": language,
            "author": item.get("author", item.get("composer", "")),
            "category": item.get("category", item.get("genre", "General")),
        })

    print(f"[VideoPsalm JSON] Found {len(songs)} songs")
    return songs


# ── VideoPsalm JSON — non-standard unquoted-key format ───────────────────────

def _parse_vp_manual(text):
    """
    Manual parser for non-standard VideoPsalm JSON (unquoted keys).
    Handles: {ID:1, Reference:"...", Verses:[{Text:"..."}]}
    """
    result = {"Songs": []}

    abbr_m = re.search(r'Abbreviation\s*:\s*"([^"]*)"', text)
    if abbr_m:
        result["Abbreviation"] = abbr_m.group(1)

    # Split into song blocks by ID:
    song_blocks = re.split(r'(?=\{?\s*ID\s*:\s*\d)', text)

    for block in song_blocks:
        song = {}

        id_m = re.search(r'ID\s*:\s*(\d+)', block)
        if id_m:
            song["ID"] = int(id_m.group(1))

        ref_m = re.search(r'Reference\s*:\s*"([^"]*)"', block)
        if ref_m:
            song["Reference"] = ref_m.group(1)

        title_m = re.search(r'Title\s*:\s*"([^"]*)"', block)
        if title_m:
            song["Title"] = title_m.group(1)

        author_m = re.search(r'Author\s*:\s*"([^"]*)"', block)
        if author_m:
            song["Author"] = author_m.group(1)

        # Extract verse texts
        verses = []
        for v in re.finditer(r'Text\s*:\s*"((?:[^"\\]|\\.)*)"', block):
            verse_text = (
                v.group(1)
                .replace("\\n", "\n")
                .replace("\\r", "")
                .replace('\\"', '"')
            )
            if verse_text.strip():
                verses.append({"Text": verse_text})

        song["Verses"] = verses
        if verses and song.get("ID"):
            result["Songs"].append(song)

    return result


def import_videopsalm_json(filepath):
    """
    Import a VideoPsalm JSON collection with potentially unquoted keys.
    Format: {Abbreviation: "AL", Songs: [{ID:1, Reference:"1", Verses:[{Text:"..."}]}]}

    Attempts:
      1. Regex-fix unquoted keys → json.loads()
      2. demjson3.decode()
      3. Manual regex parser (_parse_vp_manual)
    """
    print(f"[VP] Import: {filepath}")

    with open(filepath, "rb") as f:
        raw = f.read()

    # Strip BOM
    if raw.startswith(b"\xef\xbb\xbf"):
        raw = raw[3:]

    text = raw.decode("utf-8", errors="replace")

    data = None

    # 1. Fix unquoted keys with regex; also collapse literal newlines inside
    #    string values so json.loads doesn't choke on them.
    try:
        fixed = re.sub(r'([{,\[]\s*)([A-Za-z_]\w*)\s*:', r'\1"\2":', text)
        # Replace bare (unescaped) newlines that are inside a JSON string
        fixed = re.sub(r'(?<=[^\\])\n(?=[^"]*")', r'\\n', fixed)
        data = json.loads(fixed)
    except Exception:
        pass

    # 2. demjson3
    if data is None:
        try:
            import demjson3  # type: ignore
            data = demjson3.decode(text)
        except ImportError:
            pass
        except Exception:
            pass

    # 3. Manual parser
    if data is None:
        data = _parse_vp_manual(text)

    if not data or not data.get("Songs"):
        raise ImportError("Nu s-a putut parsa fișierul VideoPsalm JSON")

    abbreviation = data.get("Abbreviation", "")
    songs_data   = data.get("Songs", [])
    results = []

    for song in songs_data:
        verses = song.get("Verses", [])
        if not verses:
            continue

        author = song.get("Author", "") or ""

        slides = []
        for verse in verses:
            if isinstance(verse, dict):
                v_text = verse.get("Text") or verse.get("text") or ""
            else:
                v_text = str(verse)
            # Normalise literal \n sequences that survived JSON parsing
            v_text = v_text.replace("\\n", "\n").replace("\\r", "").strip()
            if v_text:
                slides.append(v_text)

        if not slides:
            continue

        # Title: prefer explicit Title/Reference field; fall back to first
        # line of the first verse text (the typical VideoPsalm convention).
        raw_title = (song.get("Title") or "").strip()
        if not raw_title:
            raw_title = (song.get("Reference") or "").strip()
        if not raw_title:
            first_line = slides[0].split("\n")[0].strip()
            raw_title = first_line or f"Cântare {song.get('ID', '')}"

        results.append({
            "title":    raw_title,
            "author":   author,
            "content":  "\n\n".join(slides),
            "slides":   slides,
            "language": "ro",
            "category": "General",
            "notes":    f"VideoPsalm {abbreviation}".strip(),
        })

    print(f"[VP] Import: {len(results)} cântări")
    return results


# ── VideoPsalm .xml ───────────────────────────────────────────────────────────

def import_videopasalm_xml(filepath_or_root, filename=""):
    if isinstance(filepath_or_root, ET.Element):
        root = filepath_or_root
    else:
        root = _try_parse_xml_with_encodings(filepath_or_root)
        if root is None:
            raise ImportError(f"Cannot parse XML: {filepath_or_root}")
        filename = str(filepath_or_root)

    songs = []
    song_nodes = (
        root.findall(".//song") or root.findall(".//Song") or [root]
    )

    for song_node in song_nodes:
        parsed = _parse_vpc_song_xml(song_node, filename)
        if parsed:
            songs.append(parsed)

    print(f"[VideoPsalm XML] Found {len(songs)} songs")
    return songs


def _parse_vpc_song_xml(song_node, filename=""):
    """
    Parse a single song XML node.
    Supports multiple VideoPsalm XML structures.
    """
    title = (
        song_node.get("title") or
        song_node.get("name") or
        _elem_text(song_node, "title") or
        _elem_text(song_node, "Title") or
        Path(filename).stem or
        "Untitled"
    )

    slides = []
    lyrics_node = song_node.find(".//lyrics") or song_node.find(".//Lyrics") or song_node

    verse_nodes = (
        lyrics_node.findall(".//verse") +
        lyrics_node.findall(".//Verse") +
        lyrics_node.findall(".//strophe") +
        lyrics_node.findall(".//stanza") +
        lyrics_node.findall(".//part")
    )

    # Remove duplicates while preserving order
    seen = set()
    unique_verses = []
    for v in verse_nodes:
        vid = id(v)
        if vid not in seen:
            seen.add(vid)
            unique_verses.append(v)

    for vn in unique_verses:
        lines_node = vn.find("lines") or vn.find("Lines")
        if lines_node is not None and lines_node.text:
            text = lines_node.text.replace("\\n", "\n").strip()
        else:
            line_nodes = vn.findall("line") + vn.findall("Line")
            if line_nodes:
                text = "\n".join(l.text.strip() for l in line_nodes if l.text and l.text.strip())
            else:
                text = ET.tostring(vn, encoding="unicode", method="text").strip()

        if text.strip():
            slides.append(text.strip())

    if not slides:
        all_text = ET.tostring(song_node, encoding="unicode", method="text").strip()
        slides = text_to_slides(all_text)

    if not slides:
        return None

    content = "\n\n".join(slides)
    return {
        "title": str(title).strip(),
        "content": content,
        "slides": slides,
        "language": song_node.get("language", song_node.get("lang", "ro")),
        "author": song_node.get("author", song_node.get("composer", _elem_text(song_node, "author") or "")),
        "category": song_node.get("category", "General"),
    }


def _elem_text(node, tag):
    """Return text of first matching child element, or None."""
    el = node.find(tag)
    if el is not None and el.text:
        return el.text.strip()
    return None


# ── VideoPsalm .vpc (ZIP of .song XML files) ─────────────────────────────────

def import_vpc(filepath):
    """
    Import a VideoPsalm .vpc file.
    .vpc is a ZIP archive containing .song files (XML) or other XML/JSON files.
    Falls back to plain XML or JSON if not a valid ZIP.
    """
    songs = []

    if zipfile.is_zipfile(filepath):
        with zipfile.ZipFile(filepath, "r") as z:
            names = z.namelist()
            print(f"[VPC] ZIP contents ({len(names)} files): {names[:20]}")

            # Priority: .song files, then .xml, then anything without a system prefix
            def priority(n):
                ext = Path(n).suffix.lower()
                if ext == ".song":
                    return 0
                if ext == ".xml":
                    return 1
                if n.startswith("__") or n.startswith("."):
                    return 99
                return 5

            candidate_files = sorted(
                [n for n in names if not n.endswith("/")],
                key=priority
            )

            for name in candidate_files:
                if priority(name) >= 99:
                    continue
                try:
                    with z.open(name) as f:
                        raw = f.read()
                    if not raw.strip():
                        continue

                    # Try XML
                    root = _try_parse_xml_with_encodings(raw)
                    if root is not None:
                        # Single song node or collection
                        if root.tag.lower() in ("song", "songs", "songbase", "songbook", "hymnal"):
                            if root.tag.lower() == "song":
                                parsed = _parse_vpc_song_xml(root, name)
                                if parsed:
                                    songs.append(parsed)
                            else:
                                for sn in root.findall(".//song") + root.findall(".//Song"):
                                    parsed = _parse_vpc_song_xml(sn, name)
                                    if parsed:
                                        songs.append(parsed)
                        else:
                            # Try as multi-song root
                            song_nodes = root.findall(".//song") + root.findall(".//Song")
                            if song_nodes:
                                for sn in song_nodes:
                                    parsed = _parse_vpc_song_xml(sn, name)
                                    if parsed:
                                        songs.append(parsed)
                            else:
                                parsed = _parse_vpc_song_xml(root, name)
                                if parsed:
                                    songs.append(parsed)
                        continue

                    # Try JSON
                    try:
                        text = raw.decode("utf-8", errors="replace")
                        data = json.loads(text)
                        items = data if isinstance(data, list) else data.get("songs", [data])
                        for item in items:
                            title = item.get("title", item.get("name", Path(name).stem))
                            verses = item.get("verses", item.get("strophes", []))
                            slides = []
                            for v in verses:
                                t = v.get("text", v.get("content", "")) if isinstance(v, dict) else v
                                if t:
                                    slides.append(str(t).strip())
                            if not slides:
                                slides = text_to_slides(item.get("content", ""))
                            if slides:
                                songs.append({
                                    "title": str(title),
                                    "content": "\n\n".join(slides),
                                    "slides": slides,
                                    "language": item.get("language", "ro"),
                                    "author": item.get("author", ""),
                                    "category": "General",
                                })
                    except (json.JSONDecodeError, UnicodeDecodeError):
                        pass

                except Exception as e:
                    print(f"[VPC] Error processing '{name}': {e}")

        print(f"[VPC] Total songs imported from ZIP: {len(songs)}")
        if songs:
            return songs

    # Not a ZIP (or ZIP was empty) — try plain XML/JSON
    print("[VPC] Not a valid ZIP or no songs found in ZIP, trying plain XML/JSON")
    try:
        return import_videopasalm_xml(filepath)
    except Exception:
        pass
    try:
        return import_videopasalm_json(filepath)
    except Exception:
        pass

    # Last resort: show first 200 bytes for diagnosis
    with open(filepath, "rb") as f:
        preview = f.read(200)
    print(f"[VPC] Could not parse. First 200 bytes: {preview!r}")
    return []


# ── EasyWorship .ewsx (ZIP archive) ──────────────────────────────────────────

def import_easyworship_ewsx(filepath):
    songs = []
    try:
        with zipfile.ZipFile(filepath, "r") as z:
            names = z.namelist()
            print(f"[EWSX] ZIP contents: {names[:20]}")
            xml_files = [n for n in names if n.lower().endswith((".xml", ".json"))]
            for name in xml_files:
                with z.open(name) as f:
                    raw = f.read()
                try:
                    data = json.loads(raw.decode("utf-8", errors="replace"))
                    items = data.get("songs", [data]) if isinstance(data, dict) else data
                    for item in items:
                        title = item.get("title", Path(name).stem)
                        slides_raw = item.get("slides", item.get("verses", []))
                        slides = [s if isinstance(s, str) else s.get("text", "") for s in slides_raw]
                        slides = [s.strip() for s in slides if s.strip()]
                        if not slides:
                            slides = text_to_slides(item.get("content", ""))
                        songs.append({
                            "title": title, "content": "\n\n".join(slides),
                            "slides": slides, "language": "ro",
                            "author": "", "category": "General",
                        })
                except (json.JSONDecodeError, Exception):
                    root = _try_parse_xml_with_encodings(raw)
                    if root is not None:
                        for sn in (root.findall(".//song") + root.findall(".//Song")) or [root]:
                            title = sn.get("title", sn.get("name", Path(name).stem))
                            slides = []
                            for slide in sn.findall(".//slide") + sn.findall(".//Slide") + sn.findall(".//verse"):
                                text = ET.tostring(slide, encoding="unicode", method="text").strip()
                                if text:
                                    slides.append(text)
                            if not slides:
                                slides = text_to_slides(
                                    ET.tostring(sn, encoding="unicode", method="text").strip()
                                )
                            songs.append({
                                "title": title, "content": "\n\n".join(slides),
                                "slides": slides, "language": "ro",
                                "author": "", "category": "General",
                            })
    except zipfile.BadZipFile:
        print("[EWSX] Not a ZIP, trying as SQLite DB")
        songs = import_easyworship_db(filepath)
    print(f"[EWSX] Found {len(songs)} songs")
    return songs


# ── EasyWorship .db (SQLite) ──────────────────────────────────────────────────

def import_easyworship_db(filepath):
    songs = []
    try:
        conn = sqlite3.connect(filepath)
        conn.row_factory = sqlite3.Row

        # List ALL tables and their columns
        tables = [
            r[0] for r in
            conn.execute("SELECT name FROM sqlite_master WHERE type='table'").fetchall()
        ]
        print(f"[EasyWorship DB] All tables: {tables}")

        for tbl in tables:
            try:
                cols = [
                    row[1] for row in
                    conn.execute(f"PRAGMA table_info({tbl})").fetchall()
                ]
                print(f"[EasyWorship DB] Table '{tbl}': columns = {cols}")
            except Exception as col_err:
                print(f"[EasyWorship DB] Cannot inspect '{tbl}': {col_err}")

        # Find the song table — EW6/7 uses "song", older versions may differ
        song_table = None
        for candidate in tables:
            cl = candidate.lower()
            if cl in ("song", "songs", "presentation", "presentations"):
                song_table = candidate
                break
        if song_table is None:
            song_table = next(
                (t for t in tables if "song" in t.lower() or "present" in t.lower()),
                None
            )

        if song_table is None:
            print("[EasyWorship DB] No song table found!")
            conn.close()
            return songs

        print(f"[EasyWorship DB] Using table: '{song_table}'")
        rows = conn.execute(f"SELECT * FROM {song_table} LIMIT 500").fetchall()
        print(f"[EasyWorship DB] Rows in '{song_table}': {len(rows)}")

        for row in rows:
            d = dict(row)

            # Title: try common column names
            title = (
                d.get("title") or d.get("Title") or d.get("name") or
                d.get("Name") or d.get("song_title") or "Untitled"
            )

            # Lyrics: try common column names
            content_raw = (
                d.get("words") or d.get("Words") or
                d.get("lyrics") or d.get("Lyrics") or
                d.get("content") or d.get("Content") or
                d.get("text") or d.get("Text") or
                d.get("SongText") or d.get("songtext") or ""
            )

            if not content_raw:
                # Try any column that looks like it could have lyrics
                for col, val in d.items():
                    if isinstance(val, str) and len(val) > 20 and col.lower() not in (
                        "title", "author", "copyright", "guid", "background",
                        "modified", "created", "category", "notes"
                    ):
                        content_raw = val
                        print(f"[EasyWorship DB] Using fallback column '{col}' for lyrics")
                        break

            if not content_raw:
                continue

            content_str = str(content_raw)

            # Strip RTF if needed
            if content_str.lstrip().startswith("{\\rtf"):
                content_str = strip_rtf(content_str)

            slides = text_to_slides(content_str)
            content = "\n\n".join(slides)

            author = str(d.get("author", d.get("Author", d.get("music_author", ""))) or "")
            songs.append({
                "title": str(title),
                "content": content,
                "slides": slides,
                "language": "ro",
                "author": author,
                "category": "General",
            })

        conn.close()
    except Exception as e:
        print(f"[EasyWorship DB] Error: {e}")

    print(f"[EasyWorship DB] Imported {len(songs)} songs")
    return songs


# ── EasyWorship 7 (.db with JSON lyrics) ─────────────────────────────────────

def detect_easyworship7_default_path():
    """Return the default EW7 song.db path if it exists on this machine."""
    docs = Path.home() / "Documents"
    candidates = [
        docs / "Easyworship" / "Default" / "Databases" / "Data" / "song.db",
        docs / "EasyWorship" / "Default" / "Databases" / "Data" / "song.db",
        Path("C:/Users") / os.environ.get("USERNAME", "") / "Documents" / "Easyworship" / "Default" / "Databases" / "Data" / "song.db",
    ]
    for p in candidates:
        if p.exists():
            return str(p)
    return None


def _parse_ew7_json_lyrics(raw) -> list[str]:
    """
    Parse EasyWorship 7 JSON lyrics field.
    Format: {"slides": [{"text": "...", "label": "Verse 1"}, ...]}
    Returns list of slide texts.
    """
    if not raw:
        return []

    # It may already be a dict (if sqlite3 decoded JSON automatically)
    if isinstance(raw, dict):
        data = raw
    else:
        raw_str = str(raw).strip()
        if not raw_str.startswith("{"):
            return []
        try:
            data = json.loads(raw_str)
        except (json.JSONDecodeError, ValueError):
            return []

    slides = data.get("slides", [])
    result = []
    for slide in slides:
        if isinstance(slide, dict):
            text = slide.get("text", slide.get("body", slide.get("content", "")))
        else:
            text = str(slide)
        text = text.strip()
        if text:
            result.append(text)
    return result


def import_easyworship7_db(filepath, progress_callback=None):
    """
    Import songs from an EasyWorship 7 song.db SQLite database.
    EW7 stores lyrics as JSON: {"slides": [{"text": "...", "label": "..."}, ...]}
    progress_callback(current, total) called periodically.
    """
    songs = []
    try:
        conn = sqlite3.connect(filepath)
        conn.row_factory = sqlite3.Row

        tables = [
            r[0] for r in
            conn.execute("SELECT name FROM sqlite_master WHERE type='table'").fetchall()
        ]
        print(f"[EW7] Tables: {tables}")

        # EW7 uses table named "song"
        song_table = next(
            (t for t in tables if t.lower() in ("song", "songs")), None
        )
        if song_table is None:
            print("[EW7] No 'song' table found, falling back to EW6 importer")
            conn.close()
            return import_easyworship_db(filepath)

        cols = [row[1] for row in conn.execute(f"PRAGMA table_info({song_table})").fetchall()]
        print(f"[EW7] Columns in '{song_table}': {cols}")

        rows = conn.execute(f"SELECT * FROM {song_table}").fetchall()
        total = len(rows)
        print(f"[EW7] Total rows: {total}")

        for i, row in enumerate(rows):
            d = dict(row)

            title = (
                d.get("title") or d.get("Title") or d.get("name") or "Untitled"
            )
            author = str(d.get("author", d.get("music_author", "")) or "")
            copyright_ = str(d.get("copyright", "") or "")

            # EW7 lyrics column is usually 'lyrics' or 'words'
            lyrics_raw = (
                d.get("lyrics") or d.get("Lyrics") or
                d.get("words") or d.get("Words") or
                d.get("content") or d.get("text") or ""
            )

            slides = []

            # Try JSON parse first (EW7 format)
            if lyrics_raw:
                lyrics_str = str(lyrics_raw).strip()
                if lyrics_str.startswith("{"):
                    slides = _parse_ew7_json_lyrics(lyrics_str)
                elif lyrics_str.lstrip().startswith("{\\rtf"):
                    stripped = strip_rtf(lyrics_str)
                    slides = text_to_slides(stripped)
                else:
                    slides = text_to_slides(lyrics_str)

            if not slides:
                continue

            content = "\n\n".join(slides)
            songs.append({
                "title": str(title).strip(),
                "content": content,
                "slides": slides,
                "language": "ro",
                "author": author.strip(),
                "category": "General",
                "notes": copyright_.strip(),
            })

            if progress_callback and i % 50 == 0:
                progress_callback(i + 1, total)

        conn.close()
        if progress_callback:
            progress_callback(total, total)

    except Exception as e:
        print(f"[EW7] Error: {e}")
        import traceback
        traceback.print_exc()

    print(f"[EW7] Imported {len(songs)} songs")
    return songs


# ── EasyWorship 7 two-file RTF layout (Songs.db + SongWords.db) ──────────────

def import_easyworship7(songs_db_path, words_db_path, progress_callback=None):
    """
    Import EasyWorship 7 split-DB layout where:
      Songs.db   → table 'song': rowid, title, author, copyright
      SongWords.db → table 'word': rowid, song_id (= rowid from Songs.db), words (RTF)

    Uses striprtf for RTF → plain text, then parses Verse/Chorus/Bridge labels.
    Falls back to import_easyworship7_twofile() if striprtf is unavailable or
    the 'word' table is not found.
    """
    try:
        from striprtf.striprtf import rtf_to_text as _rtf_to_text
    except ImportError:
        print("[EW7] striprtf not installed — falling back to twofile importer")
        return import_easyworship7_twofile(songs_db_path, words_db_path, progress_callback)

    print(f"[EW7] Songs: {songs_db_path}")
    print(f"[EW7] Words: {words_db_path}")

    LABEL_RE = re.compile(
        r'^(Verse|Chorus|Bridge|Intro|Outro|Pre-Chorus|Tag|Ending)\s*\d*\s*$',
        re.IGNORECASE,
    )

    # ── Load song metadata keyed by rowid ────────────────────────────────────
    songs_conn = sqlite3.connect(songs_db_path)
    songs_conn.row_factory = sqlite3.Row
    songs_map = {}
    try:
        for row in songs_conn.execute(
            "SELECT rowid, title, author, copyright FROM song"
        ):
            songs_map[row["rowid"]] = {
                "title":     row["title"]     or "",
                "author":    row["author"]    or "",
                "copyright": row["copyright"] or "",
            }
    except Exception as e:
        print(f"[EW7] Cannot read Songs.db: {e}")
        songs_conn.close()
        return import_easyworship7_twofile(songs_db_path, words_db_path, progress_callback)
    songs_conn.close()
    print(f"[EW7] {len(songs_map)} titles loaded")

    # ── Load RTF words keyed by song_id ──────────────────────────────────────
    words_conn = sqlite3.connect(words_db_path)
    words_conn.row_factory = sqlite3.Row

    # Check the table name — some builds use 'word', others 'SongWords'
    w_tables = [r[0] for r in words_conn.execute(
        "SELECT name FROM sqlite_master WHERE type='table'"
    ).fetchall()]
    print(f"[EW7] SongWords tables: {w_tables}")
    word_table = next(
        (t for t in w_tables if t.lower() in ("word", "words", "songwords")), None
    )
    if word_table is None:
        print("[EW7] 'word' table not found — falling back to twofile importer")
        words_conn.close()
        return import_easyworship7_twofile(songs_db_path, words_db_path, progress_callback)

    results = []
    all_rows = words_conn.execute(
        f"SELECT song_id, words FROM {word_table}"
    ).fetchall()
    total = len(all_rows)

    for i, row in enumerate(all_rows):
        song_id = row["song_id"]
        rtf     = row["words"] or ""
        meta    = songs_map.get(song_id, {})
        title   = meta.get("title", "")
        if not title:
            continue

        # Strip RTF → plain text
        try:
            text = _rtf_to_text(rtf).strip()
        except Exception:
            text = strip_rtf(rtf)   # fallback to built-in

        if not text:
            continue

        # Parse slides by section label
        slides = []
        current_lines = []

        for line in text.splitlines():
            stripped = line.strip()
            if LABEL_RE.match(stripped):
                if current_lines:
                    chunk = "\n".join(current_lines).strip()
                    if chunk:
                        slides.append(chunk)
                    current_lines = []
            elif not stripped:
                if current_lines:
                    chunk = "\n".join(current_lines).strip()
                    if chunk:
                        slides.append(chunk)
                    current_lines = []
            else:
                current_lines.append(stripped)

        if current_lines:
            chunk = "\n".join(current_lines).strip()
            if chunk:
                slides.append(chunk)

        slides = [s for s in slides if len(s.strip()) > 2]

        if not slides:
            # Fallback: chunk by 4 lines
            all_lines = [l for l in text.splitlines() if l.strip()]
            for idx in range(0, len(all_lines), 4):
                chunk = all_lines[idx:idx + 4]
                if chunk:
                    slides.append("\n".join(chunk))

        if not slides:
            continue

        results.append({
            "title":    title,
            "author":   meta.get("author", ""),
            "content":  "\n\n".join(slides),
            "slides":   slides,
            "language": "ro",
            "category": "General",
            "notes":    meta.get("copyright", ""),
        })

        if progress_callback and i % 50 == 0:
            progress_callback(i + 1, total)

    words_conn.close()
    if progress_callback:
        progress_callback(total, total)

    print(f"[EW7] Import: {len(results)} cântări")
    return results


# ── EasyWorship 7 two-file layout (Songs.db + SongsWords.db) ─────────────────

def import_easyworship7_twofile(songs_db_path, words_db_path, progress_callback=None):
    """
    Import songs from EasyWorship 7's split-database layout.

    songs_db_path  – path to Songs.db   (table 'song': uid, title, author, copyright)
    words_db_path  – path to SongsWords.db (table 'song_slide': uid FK, ordinal, RTF content)

    Returns a list of song dicts compatible with db.add_song().
    """
    songs = []
    try:
        conn_s = sqlite3.connect(songs_db_path)
        conn_s.row_factory = sqlite3.Row
        conn_w = sqlite3.connect(words_db_path)
        conn_w.row_factory = sqlite3.Row

        # ── Discover tables ──────────────────────────────────────────────────
        s_tables = [r[0] for r in conn_s.execute(
            "SELECT name FROM sqlite_master WHERE type='table'"
        ).fetchall()]
        w_tables = [r[0] for r in conn_w.execute(
            "SELECT name FROM sqlite_master WHERE type='table'"
        ).fetchall()]
        print(f"[EW7-2F] Songs.db tables:      {s_tables}")
        print(f"[EW7-2F] SongsWords.db tables: {w_tables}")

        # ── Song metadata table (Songs.db) ───────────────────────────────────
        song_table = next(
            (t for t in s_tables if t.lower() in ("song", "songs")), None
        )
        if song_table is None:
            song_table = next(
                (t for t in s_tables if "song" in t.lower()), None
            )
        if song_table is None:
            print("[EW7-2F] No song table found in Songs.db — aborting")
            conn_s.close(); conn_w.close()
            return songs

        s_cols = [r[1] for r in conn_s.execute(f"PRAGMA table_info({song_table})").fetchall()]
        print(f"[EW7-2F] '{song_table}' columns: {s_cols}")

        song_rows = conn_s.execute(f"SELECT * FROM {song_table}").fetchall()
        total = len(song_rows)
        print(f"[EW7-2F] Total songs in metadata table: {total}")

        # uid → {title, author, copyright}
        meta_map = {}
        for row in song_rows:
            d = dict(row)
            uid = (d.get("song_uid") or d.get("uid") or
                   d.get("SongUID") or d.get("id") or d.get("rowid"))
            title = (d.get("title") or d.get("Title") or
                     d.get("name") or "Untitled")
            author = str(d.get("author") or d.get("music_author") or "")
            copyright_ = str(d.get("copyright") or "")
            if uid is not None:
                meta_map[uid] = {
                    "title": str(title).strip(),
                    "author": author.strip(),
                    "copyright": copyright_.strip(),
                }

        # ── Slide table (SongsWords.db) ──────────────────────────────────────
        slide_table = next(
            (t for t in w_tables
             if t.lower() in ("song_slide", "songslide", "slide", "slides")),
            None,
        )
        if slide_table is None:
            slide_table = next(
                (t for t in w_tables if "slide" in t.lower() or "word" in t.lower()),
                None,
            )

        slides_by_uid: dict = {}   # uid → [(ordinal, text), …]
        if slide_table is not None:
            w_cols = [r[1] for r in conn_w.execute(
                f"PRAGMA table_info({slide_table})"
            ).fetchall()]
            print(f"[EW7-2F] '{slide_table}' columns: {w_cols}")

            slide_rows = conn_w.execute(
                f"SELECT * FROM {slide_table} ORDER BY rowid"
            ).fetchall()
            for row in slide_rows:
                d = dict(row)
                uid = (d.get("song_uid") or d.get("SongUID") or
                       d.get("uid") or d.get("song_id"))
                ordinal = int(
                    d.get("slide_number") or d.get("ordinal") or
                    d.get("position") or d.get("order") or 0
                )
                raw = (
                    d.get("slide_content") or d.get("content") or
                    d.get("lyrics") or d.get("words") or
                    d.get("text") or ""
                )
                text = str(raw).strip() if raw else ""
                # Strip RTF if present
                if text.lstrip().startswith("{\\rtf"):
                    text = strip_rtf(text)
                if text and uid is not None:
                    slides_by_uid.setdefault(uid, []).append((ordinal, text))
        else:
            print("[EW7-2F] No slide table found in SongsWords.db")

        conn_s.close()
        conn_w.close()

        # ── Assemble song dicts ──────────────────────────────────────────────
        for i, (uid, meta) in enumerate(meta_map.items()):
            raw_slides = slides_by_uid.get(uid, [])
            raw_slides.sort(key=lambda x: x[0])          # sort by ordinal
            slide_texts = [s for _, s in raw_slides if s]

            if not slide_texts:
                # Slides DB had no rows for this UID — skip gracefully
                continue

            content = "\n\n".join(slide_texts)
            songs.append({
                "title":    meta["title"],
                "content":  content,
                "slides":   slide_texts,
                "language": "ro",
                "author":   meta["author"],
                "category": "General",
                "notes":    meta["copyright"],
            })

            if progress_callback and i % 50 == 0:
                progress_callback(i + 1, total)

        if progress_callback:
            progress_callback(total, total)

    except Exception as e:
        print(f"[EW7-2F] Error: {e}")
        import traceback
        traceback.print_exc()

    print(f"[EW7-2F] Imported {len(songs)} songs")
    return songs


# ── BibleShow .bib ────────────────────────────────────────────────────────────

# ── Microsoft Access DB detection ─────────────────────────────────────────────

def _is_access_db(filepath) -> bool:
    """Return True if the file header matches a Microsoft Access / Jet DB."""
    try:
        with open(filepath, "rb") as f:
            header = f.read(32)
        return b"Standard Jet DB" in header or b"Standard ACE DB" in header
    except Exception:
        return False


# ── Access DB shared helpers ───────────────────────────────────────────────────

def _access_find_tables(tables):
    """
    Given a list of table name strings, return (book_table, verse_table) using
    name heuristics.  Either value may be None if no good match is found.
    """
    book_keywords  = ['book', 'carte', 'bibbook', 'biblebok', 'buch']
    verse_keywords = ['verse', 'verset', 'text', 'scripture', 'bibeltext',
                      'inhalt', 'content', 'vers']

    book_table = next(
        (t for t in tables if any(k in t.lower() for k in book_keywords)), None
    )
    verse_table = next(
        (t for t in tables if any(k in t.lower() for k in verse_keywords)
         and t != book_table),
        None,
    )
    return book_table, verse_table


def _access_parse_books(rows_dicts):
    """
    Convert a list of column→value dicts (lowercase keys) to book records.
    """
    books = []
    for i, d in enumerate(rows_dicts, 1):
        name = (
            d.get('longname') or d.get('long_name') or
            d.get('name') or d.get('bookname') or
            d.get('bname') or d.get('title') or f"Book{i}"
        )
        abbr = (
            d.get('shortname') or d.get('short_name') or
            d.get('abbr') or d.get('bsname') or str(name)[:3]
        )
        books.append({
            "id": i,
            "name": str(name).strip(),
            "abbreviation": str(abbr).strip(),
            "testament": "OT" if i <= 39 else "NT",
            "book_order": i,
        })
    return books


def _access_parse_verses(rows_dicts):
    """
    Convert a list of column→value dicts (lowercase keys) to verse records.
    """
    verses = []
    for d in rows_dicts:
        book_id   = d.get('book') or d.get('bookid') or d.get('book_id') or d.get('b') or 1
        chapter   = d.get('chapter') or d.get('chap') or d.get('c') or 1
        verse_num = d.get('verse') or d.get('vers') or d.get('v') or 1
        text      = (
            d.get('text') or d.get('scripture') or d.get('versetext') or
            d.get('inhalt') or d.get('content') or ''
        )
        try:
            verses.append({
                "book_id": int(book_id),
                "chapter": int(chapter),
                "verse":   int(verse_num),
                "text":    str(text).strip(),
            })
        except (TypeError, ValueError):
            continue
    return verses


# ── Method 1: mdbtools (subprocess) ───────────────────────────────────────────

def ensure_mdbtools():
    """
    Localizează executabilele mdbtools (mdb-tables / mdb-export).

    Ordinea de căutare:
      1. PATH — mdb-tables disponibil direct în shell
      2. <app_dir>/mdbtools/     — executabile incluse manual sau descărcate
      3. Directoare fixe uzuale  — C:\\mdbtools, C:\\tools\\mdbtools, etc.
      4. <python>/Scripts/       — pip install mdbtools-win
      5. Descărcare automată din GitHub (lsgunth/mdbtools-win)

    Returnează (True, cale_mdb_tables) sau (False, None).
    """
    import subprocess

    # 1. Verifică PATH
    try:
        subprocess.run(
            ["mdb-tables", "--help"],
            capture_output=True,
            timeout=5,
        )
        return True, "mdb-tables"
    except FileNotFoundError:
        pass

    # 2 & 3. Caută în directoare cunoscute
    app_dir = os.path.dirname(os.path.abspath(__file__))
    import sys
    search_dirs = [
        os.path.join(app_dir, "mdbtools"),
        os.path.join(app_dir, "bin"),
        r"C:\mdbtools",
        r"C:\mdb-tools",
        r"C:\mdbtools-app",
        r"C:\Program Files\mdbtools",
        r"C:\tools\mdbtools",
        os.path.join(os.path.dirname(sys.executable), "Scripts"),
    ]
    for d in search_dirs:
        exe = os.path.join(d, "mdb-tables.exe")
        if os.path.isfile(exe):
            print(f"[BIB] mdbtools găsit: {exe}")
            return True, exe

    # 4. Descărcare automată din GitHub
    print("[BIB] mdbtools lipsește — încerc descărcarea automată…")
    try:
        from download_mdbtools import download_mdbtools
        mdb_dir = download_mdbtools()
        if mdb_dir:
            exe = os.path.join(mdb_dir, "mdb-tables.exe")
            if os.path.isfile(exe):
                return True, exe
    except Exception as e:
        print(f"[BIB] Auto-download eșuat: {e}")

    return False, None


def run_mdb_export(mdb_export_exe, filepath, table, env, timeout=120):
    """
    Rulează mdb-export și returnează textul CSV cu encoding auto-detectat.
    Încearcă utf-8, cp1250, iso-8859-2, latin-1 în ordine;
    returnează prima variantă fără caractere de înlocuire (U+FFFD).
    """
    import subprocess
    r = subprocess.run(
        [mdb_export_exe, filepath, table],
        capture_output=True,
        timeout=timeout,
        env=env,
    )
    raw = r.stdout
    for enc in ["utf-8", "cp1250", "iso-8859-2", "latin-1"]:
        try:
            text = raw.decode(enc)
            if "�" not in text[:2000]:
                print(f"[BIB] Encoding detectat: {enc}")
                return text
        except (UnicodeDecodeError, Exception):
            continue
    return raw.decode("latin-1", errors="replace")


def import_bib_mdbtools(filepath):
    """
    Importă un fișier BibleShow .bib (Access JET/ACE) folosind mdbtools CLI.

    Structura BibleShow:
      • Structure — BibPosition, DisplayTitle, ListTitle, Abbreviation  (cărți)
      • Bible     — Book, Chapter, Verse, Scripture                     (versete)

    Dacă tabelele exacte lipsesc, cade înapoi pe euristici generice.
    Ridică ImportError dacă mdbtools nu poate fi găsit/descărcat.
    """
    import subprocess
    import csv
    import io

    # ── Localizează mdbtools ──────────────────────────────────────────────────
    ok, mdb_tables_exe = ensure_mdbtools()
    if not ok:
        show_mdbtools_help()
        raise ImportError(
            "mdbtools nu este instalat.\n"
            "Descarcă executabilele și plasează-le în:\n"
            f"  {os.path.join(os.path.dirname(os.path.abspath(__file__)), 'mdbtools')}\\"
        )

    mdb_export_exe = mdb_tables_exe.replace("mdb-tables", "mdb-export")
    mdb_dir = os.path.dirname(os.path.abspath(mdb_tables_exe))

    # Adaugă folderul mdbtools la PATH al procesului copil
    # (DLL-urile dependente trebuie să fie găsite)
    env = os.environ.copy()
    if mdb_dir and mdb_dir != ".":
        env["PATH"] = mdb_dir + os.pathsep + env.get("PATH", "")

    print(f"[BIB] mdb-tables: {mdb_tables_exe}")

    # ── Listează tabelele ────────────────────────────────────────────────────
    r = subprocess.run(
        [mdb_tables_exe, filepath],
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
        timeout=30,
        env=env,
    )
    if r.returncode != 0:
        raise ImportError(f"mdb-tables a returnat eroare: {r.stderr.strip()}")

    tables = r.stdout.strip().split()
    if not tables:
        raise ImportError("mdbtools: niciun tabel găsit în fișier")
    print(f"[BIB] Tabele: {tables}")

    # ── Helper: exportă un tabel ca listă de rânduri (encoding auto-detectat) ──
    def _export(tbl, timeout=120):
        text = run_mdb_export(mdb_export_exe, filepath, tbl, env, timeout)
        reader = csv.DictReader(io.StringIO(text))
        return list(reader)

    # ── Cărți din tabela Structure (BibleShow) ───────────────────────────────
    books = []
    if "Structure" in tables:
        print("[BIB] Extrag cărți din 'Structure'")
        for row in _export("Structure"):
            try:
                bib_pos = int(row.get("BibPosition") or 0)
            except (ValueError, TypeError):
                continue
            if not bib_pos:
                continue
            name = (
                row.get("DisplayTitle") or
                row.get("ListTitle") or
                f"Carte{bib_pos}"
            ).strip()
            abbr = (row.get("Abbreviation") or name[:3]).strip().rstrip(".")
            books.append({
                "id":           bib_pos,
                "name":         name,
                "abbreviation": abbr,
                "testament":    "OT" if bib_pos <= 39 else "NT",
                "book_order":   bib_pos,
            })
        books.sort(key=lambda x: x["id"])
        print(f"[BIB] Cărți: {len(books)}")
    else:
        # Fallback generic
        book_table, _ = _access_find_tables(tables)
        if book_table:
            print(f"[BIB] Fallback cărți: '{book_table}'")
            books = _access_parse_books(
                [{k.lower(): v for k, v in row.items()}
                 for row in _export(book_table)]
            )

    # ── Versete din tabela Bible (BibleShow) ─────────────────────────────────
    verses = []
    if "Bible" in tables:
        print("[BIB] Extrag versete din 'Bible'")
        for row in _export("Bible", timeout=180):
            try:
                book_id   = int(row.get("Book")    or 0)
                chapter   = int(row.get("Chapter") or 0)
                verse_num = int(row.get("Verse")   or 0)
                text      = (row.get("Scripture")  or "").strip()
            except (ValueError, TypeError):
                continue
            if not (book_id and chapter and verse_num and text):
                continue
            # Curăță marcaje BibleShow: ^ (Strong) și taguri <…>
            text = re.sub(r"\^", "", text)
            text = re.sub(r"<[^>]+>", "", text)
            text = text.strip()
            if not text:
                continue
            verses.append({
                "book_id": book_id,
                "chapter": chapter,
                "verse":   verse_num,
                "text":    text,
            })
        print(f"[BIB] Versete: {len(verses)}")
    else:
        # Fallback generic
        _, verse_table = _access_find_tables(tables)
        if verse_table:
            print(f"[BIB] Fallback versete: '{verse_table}'")
            verses = _access_parse_verses(
                [{k.lower(): v for k, v in row.items()}
                 for row in _export(verse_table, timeout=180)]
            )

    return {"books": books, "verses": verses}


def show_mdbtools_help(parent=None):
    """
    Afișează un dialog cu instrucțiuni clare pentru instalarea mdbtools.
    Butonul „Deschide pagina de download" deschide GitHub Releases în browser.
    """
    app_dir = os.path.dirname(os.path.abspath(__file__))
    mdb_dir = os.path.join(app_dir, "mdbtools")
    try:
        from PyQt6.QtWidgets import QMessageBox
        msg = QMessageBox(parent)
        msg.setWindowTitle("Import Biblie — Setup necesar")
        msg.setText(
            "Pentru a importa fișiere .bib din BibleShow este necesar\n"
            "programul auxiliar mdbtools.\n\n"
            "Pași de instalare:\n\n"
            "1. Deschide linkul de mai jos\n"
            "2. Descarcă fișierul .zip din secțiunea Assets\n"
            "3. Extrage fișierele .exe și .dll în:\n"
            f"   {mdb_dir}\\\n\n"
            "4. Încearcă din nou importul"
        )
        btn_dl = msg.addButton(
            "Deschide pagina de download",
            QMessageBox.ButtonRole.ActionRole,
        )
        msg.addButton(QMessageBox.StandardButton.Ok)
        msg.exec()
        if msg.clickedButton() == btn_dl:
            import webbrowser
            webbrowser.open(
                "https://github.com/lsgunth/mdbtools-win/releases"
            )
    except Exception:
        # Fallback non-GUI (context fără Qt sau thread secundar)
        print(
            "[BIB] mdbtools lipsește.\n"
            f"Descarcă din: https://github.com/lsgunth/mdbtools-win/releases\n"
            f"Plasează .exe și .dll în: {mdb_dir}"
        )


# ── Method 2: pyodbc (multiple connection string variants) ────────────────────

def _import_bib_pyodbc(filepath):
    """
    Try four pyodbc connection strings in order, stopping at the first that works.
    Raises ImportError if pyodbc is not installed or all variants fail.
    """
    try:
        import pyodbc
    except ImportError:
        raise ImportError("pyodbc not installed. Run: pip install pyodbc")

    conn_strings = [
        # Variant 1 — standard ACE/Jet driver
        (
            r'DRIVER={Microsoft Access Driver (*.mdb, *.accdb)};'
            f'DBQ={filepath};'
        ),
        # Variant 2 — add explicit credentials (suppresses some registry prompts)
        (
            r'DRIVER={Microsoft Access Driver (*.mdb, *.accdb)};'
            f'DBQ={filepath};Uid=Admin;Pwd=;'
        ),
        # Variant 3 — read-only, exclusive=0 (avoids temp DSN registry write)
        (
            r'DRIVER={Microsoft Access Driver (*.mdb, *.accdb)};'
            f'DBQ={filepath};Exclusive=0;ReadOnly=1;'
        ),
        # Variant 4 — legacy *.mdb-only driver
        (
            r'DRIVER={Microsoft Access Driver (*.mdb)};'
            f'DBQ={filepath};'
        ),
    ]

    last_err = None
    for i, conn_str in enumerate(conn_strings, 1):
        try:
            conn   = pyodbc.connect(conn_str, timeout=10)
            cursor = conn.cursor()
            tables = [row.table_name for row in cursor.tables(tableType='TABLE')]
            print(f"[BIB] pyodbc variant {i} connected — tables: {tables}")

            book_table, verse_table = _access_find_tables(tables)

            # Row-count fallback
            if not book_table or not verse_table:
                counts = {}
                for t in tables:
                    try:
                        cursor.execute(f"SELECT COUNT(*) FROM [{t}]")
                        counts[t] = cursor.fetchone()[0]
                    except Exception:
                        counts[t] = 0
                sorted_t = sorted(counts, key=counts.get)
                if not book_table and sorted_t:
                    book_table = sorted_t[0]
                remaining = [t for t in sorted_t if t != book_table]
                if not verse_table and remaining:
                    verse_table = remaining[-1]

            print(f"[BIB] pyodbc — book: {book_table}  verse: {verse_table}")

            def _fetch(tbl):
                try:
                    cursor.execute(
                        f"SELECT * FROM [{tbl}] ORDER BY book, chapter, verse"
                    )
                except Exception:
                    cursor.execute(f"SELECT * FROM [{tbl}]")
                cols = [c[0].lower() for c in cursor.description]
                return [dict(zip(cols, row)) for row in cursor.fetchall()]

            books  = _access_parse_books(_fetch(book_table))   if book_table  else []
            verses = _access_parse_verses(_fetch(verse_table)) if verse_table else []
            conn.close()
            return {"books": books, "verses": verses}

        except Exception as e:
            last_err = e
            print(f"[BIB] pyodbc variant {i} failed: {e}")
            continue

    raise ImportError(f"pyodbc: all connection variants failed. Last error: {last_err}")


# ── Method 3: pure-Python Access reader libraries ─────────────────────────────

def _import_bib_accdb_lib(filepath):
    """
    Try several pure-Python Access-reading libraries:
      accdb, mdb_reader, accdbtools
    Raises ImportError if none are installed or readable.
    """
    # --- accdb ---
    try:
        import accdb  # type: ignore
        db_obj = accdb.open(filepath)
        tables = list(db_obj.table_names())
        print(f"[BIB] accdb tables: {tables}")

        book_table, verse_table = _access_find_tables(tables)

        def _rows(tbl):
            return [{k.lower(): v for k, v in row.items()}
                    for row in db_obj.table(tbl)]

        books  = _access_parse_books(_rows(book_table))   if book_table  else []
        verses = _access_parse_verses(_rows(verse_table)) if verse_table else []
        return {"books": books, "verses": verses}
    except ImportError:
        pass
    except Exception as e:
        print(f"[BIB] accdb library error: {e}")

    # --- mdb_reader ---
    try:
        import mdb_reader  # type: ignore
        with open(filepath, "rb") as fh:
            db_obj = mdb_reader.Database(fh)
        tables = list(db_obj.table_names())
        print(f"[BIB] mdb_reader tables: {tables}")

        book_table, verse_table = _access_find_tables(tables)

        def _rows_mdb(tbl):
            t = db_obj.table(tbl)
            cols = [c.name.lower() for c in t.columns]
            return [dict(zip(cols, row)) for row in t.rows()]

        books  = _access_parse_books(_rows_mdb(book_table))   if book_table  else []
        verses = _access_parse_verses(_rows_mdb(verse_table)) if verse_table else []
        return {"books": books, "verses": verses}
    except ImportError:
        pass
    except Exception as e:
        print(f"[BIB] mdb_reader library error: {e}")

    # --- accdbtools ---
    try:
        import accdbtools  # type: ignore
        db_obj = accdbtools.open_database(filepath)
        tables = db_obj.list_tables()
        print(f"[BIB] accdbtools tables: {tables}")

        book_table, verse_table = _access_find_tables(tables)

        def _rows_at(tbl):
            return [{k.lower(): v for k, v in row.items()}
                    for row in db_obj.read_table(tbl)]

        books  = _access_parse_books(_rows_at(book_table))   if book_table  else []
        verses = _access_parse_verses(_rows_at(verse_table)) if verse_table else []
        return {"books": books, "verses": verses}
    except ImportError:
        pass
    except Exception as e:
        print(f"[BIB] accdbtools library error: {e}")

    raise ImportError(
        "No pure-Python Access library available. "
        "Try: pip install accdb  or  pip install mdb-reader"
    )


# ── Public entry point ─────────────────────────────────────────────────────────

def import_bib_access(filepath):
    """
    Import a BibleShow .bib file that is a Microsoft Access (JET/ACE) database.

    Tries three methods in order, using the first that succeeds:
      1. mdbtools CLI (mdb-tables / mdb-export)
      2. pyodbc with multiple connection string variants
      3. Pure-Python accdb / mdb_reader / accdbtools libraries

    Prints which method succeeded and the book/verse counts.
    """
    methods = [
        ("mdbtools",  import_bib_mdbtools),
        ("pyodbc",    _import_bib_pyodbc),
        ("accdb-lib", _import_bib_accdb_lib),
    ]
    errors = []
    for name, fn in methods:
        try:
            result = fn(filepath)
            if result.get("books"):
                print(
                    f"[BIB] ✓ Access import via {name}: "
                    f"{len(result['books'])} cărți, {len(result['verses'])} versete"
                )
                return result
            # Method ran but returned nothing useful — keep trying
            print(f"[BIB] {name} returned no books, trying next method")
        except ImportError as e:
            errors.append(f"{name}: {e}")
            print(f"[BIB] {name} not available — {e}")
        except Exception as e:
            errors.append(f"{name}: {e}")
            print(f"[BIB] {name} error — {e}")

    raise ImportError(
        "Nu s-a putut citi fișierul .bib (Access DB).\n"
        "Încearcă: pip install pyodbc  sau  pip install accdb\n"
        "Erori detaliate:\n" + "\n".join(f"  • {e}" for e in errors)
    )


def import_bib(filepath):
    """
    Parse BibleShow .bib format.
    Order of attempts:
      0. Microsoft Access (JET/ACE) — tried first when header matches
      1. XML with encoding detection
      2. SQLite
      3. Plain text
    """
    # 0. Microsoft Access DB (JET/ACE) — check header signature first
    if _is_access_db(filepath):
        print(f"[BIB] Detected Microsoft Access DB header: {filepath}")
        try:
            result = import_bib_access(filepath)
            if result["books"]:
                return result
            print("[BIB] Access import returned no books, trying fallbacks")
        except ImportError as e:
            print(f"[BIB] {e}")
        except Exception as e:
            print(f"[BIB] Access import failed: {e}, trying fallbacks")

    # 1. Try XML with encoding detection
    print(f"[BIB] Trying XML parse: {filepath}")
    root = _try_parse_xml_with_encodings(filepath)
    if root is not None:
        result = _parse_bib_xml(root)
        if result["books"]:
            print(f"[BIB] XML: {len(result['books'])} books, {len(result['verses'])} verses")
            return result

    # 2. Try SQLite
    print("[BIB] Trying SQLite parse")
    try:
        conn = sqlite3.connect(filepath)
        tables = [r[0] for r in conn.execute(
            "SELECT name FROM sqlite_master WHERE type='table'"
        ).fetchall()]
        print(f"[BIB] SQLite tables: {tables}")
        result = _parse_bib_sqlite(conn)
        conn.close()
        if result["books"]:
            print(f"[BIB] SQLite: {len(result['books'])} books, {len(result['verses'])} verses")
            return result
    except Exception as e:
        print(f"[BIB] SQLite failed: {e}")

    # 3. Try plain text with auto-detected encoding
    print("[BIB] Trying plain text parse")
    enc = detect_encoding(filepath)
    print(f"[BIB] Detected encoding: {enc}")

    for try_enc in [enc, "utf-16", "utf-16-le", "utf-16-be",
                    "utf-8-sig", "utf-8", "cp1250", "cp1252", "iso-8859-2", "latin-1"]:
        try:
            with open(filepath, "r", encoding=try_enc, errors="replace") as f:
                content = f.read()

            if _has_garbage(content):
                print(f"[BIB] Encoding {try_enc} produced garbage, skipping")
                continue

            result = _parse_bib_text(content)
            if result["books"]:
                print(f"[BIB] Plain text ({try_enc}): {len(result['books'])} books, {len(result['verses'])} verses")
                return result
        except (UnicodeDecodeError, UnicodeError):
            continue

    # 4. Raw-bytes diagnosis
    with open(filepath, "rb") as f:
        raw = f.read(200)
    print(f"[BIB] Cannot parse. First 200 bytes: {raw!r}")
    raise ImportError(f"Cannot parse .bib file. Check console output for details.")


def _parse_bib_xml(root):
    books = []
    verses = []
    book_order = 1

    book_nodes = (
        root.findall(".//b") or root.findall(".//book") or root.findall(".//Book") or
        root.findall(".//BIBLEBOOK") or root.findall(".//biblebook")
    )

    for bn in book_nodes:
        book_name = (
            bn.get("n") or bn.get("name") or bn.get("title") or
            bn.get("bnlong") or bn.get("bname") or f"Book{book_order}"
        )
        book_abbr = bn.get("abbr") or bn.get("short") or bn.get("bsname") or book_name[:3]
        testament = "OT" if book_order <= 39 else "NT"
        book_id = book_order

        books.append({
            "id": book_id, "name": book_name, "abbreviation": book_abbr,
            "testament": testament, "book_order": book_order
        })

        chapter_nodes = (
            bn.findall(".//c") or bn.findall(".//chapter") or
            bn.findall(".//CHAPTER") or bn.findall(".//chapter")
        )
        for cn in chapter_nodes:
            chap_num = int(cn.get("n") or cn.get("number") or cn.get("cnumber") or 1)
            verse_nodes = cn.findall(".//v") or cn.findall(".//verse") or cn.findall(".//VERS")
            for vn in verse_nodes:
                verse_num = int(vn.get("n") or vn.get("number") or vn.get("vnumber") or 1)
                text = vn.text or vn.get("text") or ""
                # Strip any embedded XML tags from verse text
                if "<" in text:
                    text = ET.tostring(vn, encoding="unicode", method="text")
                verses.append({
                    "book_id": book_id, "chapter": chap_num,
                    "verse": verse_num, "text": text.strip()
                })

        book_order += 1

    return {"books": books, "verses": verses}


def _parse_bib_sqlite(conn):
    conn.row_factory = sqlite3.Row
    tables = [
        r[0] for r in
        conn.execute("SELECT name FROM sqlite_master WHERE type='table'").fetchall()
    ]
    books = []
    verses = []

    book_table = next((t for t in tables if "book" in t.lower()), None)
    verse_table = next(
        (t for t in tables if "verse" in t.lower() or "text" in t.lower()), None
    )

    if book_table:
        for i, row in enumerate(
            conn.execute(f"SELECT * FROM {book_table} ORDER BY rowid").fetchall(), 1
        ):
            d = dict(row)
            name = (
                d.get("name") or d.get("long_name") or d.get("title") or
                d.get("book_name") or f"Book{i}"
            )
            abbr = d.get("short_name") or d.get("abbr") or d.get("abbrev") or name[:3]
            books.append({
                "id": i, "name": name, "abbreviation": abbr,
                "testament": "OT" if i <= 39 else "NT", "book_order": i
            })

    if verse_table:
        for row in conn.execute(f"SELECT * FROM {verse_table}").fetchall():
            d = dict(row)
            text = d.get("text") or d.get("scripture") or d.get("verse_text") or ""
            if _has_garbage(str(text)):
                # Try decoding as bytes if text is mangled
                pass
            verses.append({
                "book_id": d.get("book_id") or d.get("book") or 1,
                "chapter": d.get("chapter") or d.get("chap") or 1,
                "verse": d.get("verse") or d.get("vers") or 1,
                "text": str(text),
            })

    return {"books": books, "verses": verses}


def _parse_bib_text(content):
    """Parse plain-text Bible: 'BookName Chapter:Verse Text' per line."""
    books = {}
    verses = []
    book_order = 1

    for line in content.splitlines():
        line = line.strip()
        if not line:
            continue

        # Split into reference and text
        parts = line.split(" ", 1)
        if len(parts) < 2:
            continue
        ref, text = parts[0], parts[1]
        if ":" not in ref:
            # Try "Book Chapter Verse text" style (space-separated)
            parts2 = line.split()
            if len(parts2) >= 4:
                try:
                    verse_num = int(parts2[-1])
                    chapter = int(parts2[-2])
                    book_name = " ".join(parts2[:-2])
                    text = ""
                    # This style doesn't have inline text — skip
                except ValueError:
                    pass
            continue

        book_chap, verse_str = ref.rsplit(":", 1)
        tokens = book_chap.rsplit(" ", 1)
        if len(tokens) == 2:
            book_name, chapter = tokens
        else:
            book_name, chapter = tokens[0], "1"

        try:
            chapter = int(chapter)
            verse_num = int(verse_str)
        except ValueError:
            continue

        if book_name not in books:
            books[book_name] = book_order
            book_order += 1

        book_id = books[book_name]
        verses.append({
            "book_id": book_id, "chapter": chapter,
            "verse": verse_num, "text": text.strip()
        })

    books_list = [
        {
            "id": bid, "name": bname, "abbreviation": bname[:3],
            "testament": "OT" if bid <= 39 else "NT", "book_order": bid
        }
        for bname, bid in books.items()
    ]
    return {"books": books_list, "verses": verses}


# ── Universal dispatcher ──────────────────────────────────────────────────────

def import_file(filepath):
    """Auto-detect format and import. Returns {"type": "songs"/"bible", "data": ...}"""
    ext = Path(filepath).suffix.lower()
    size = os.path.getsize(filepath)
    print(f"\n[import_file] File: {filepath}")
    print(f"[import_file] Extension: {ext}  Size: {size} bytes")

    if ext == ".txt":
        data = import_txt(filepath)
        print(f"[import_file] Result: {len(data)} songs")
        return {"type": "songs", "data": data}

    elif ext == ".docx":
        data = import_docx(filepath)
        print(f"[import_file] Result: {len(data)} songs")
        return {"type": "songs", "data": data}

    elif ext == ".pdf":
        data = import_pdf(filepath)
        print(f"[import_file] Result: {len(data)} songs")
        return {"type": "songs", "data": data}

    elif ext == ".json":
        # Try new VP JSON (unquoted-key format) first, then legacy parser
        try:
            data = import_videopsalm_json(filepath)
            if data:
                print(f"[import_file] Result: {len(data)} songs (VP JSON)")
                return {"type": "songs", "data": data}
        except Exception:
            pass
        data = import_videopasalm_json(filepath)
        print(f"[import_file] Result: {len(data)} songs")
        return {"type": "songs", "data": data}

    elif ext == ".xml":
        data = import_videopasalm_xml(filepath)
        print(f"[import_file] Result: {len(data)} songs")
        return {"type": "songs", "data": data}

    elif ext == ".vpc":
        data = import_vpc(filepath)
        if not data:
            with open(filepath, "rb") as f:
                preview = f.read(200)
            print(f"[import_file] WARNING: 0 songs from .vpc. First 200 bytes: {preview!r}")
        else:
            print(f"[import_file] Result: {len(data)} songs")
        return {"type": "songs", "data": data}

    elif ext in (".ewsx", ".db"):
        if zipfile.is_zipfile(filepath):
            data = import_easyworship_ewsx(filepath)
        else:
            data = import_easyworship_db(filepath)
        if not data:
            with open(filepath, "rb") as f:
                preview = f.read(200)
            print(f"[import_file] WARNING: 0 songs from EasyWorship. First 200 bytes: {preview!r}")
        else:
            print(f"[import_file] Result: {len(data)} songs")
        return {"type": "songs", "data": data}

    elif ext == ".bib":
        data = import_bib(filepath)
        print(f"[import_file] Result: {len(data['books'])} books, {len(data['verses'])} verses")
        return {"type": "bible", "data": data}

    else:
        # Unknown extension — try heuristics
        print(f"[import_file] Unknown extension '{ext}', trying heuristics")
        if zipfile.is_zipfile(filepath):
            print("[import_file] Detected ZIP, trying as .vpc")
            data = import_vpc(filepath)
            if data:
                return {"type": "songs", "data": data}
        try:
            conn = sqlite3.connect(filepath)
            tables = [r[0] for r in conn.execute(
                "SELECT name FROM sqlite_master WHERE type='table'"
            ).fetchall()]
            conn.close()
            if tables:
                print(f"[import_file] Detected SQLite with tables: {tables}")
                if any("song" in t.lower() for t in tables):
                    return {"type": "songs", "data": import_easyworship_db(filepath)}
                if any("book" in t.lower() or "verse" in t.lower() for t in tables):
                    return {"type": "bible", "data": import_bib(filepath)}
        except Exception:
            pass
        raise ValueError(f"Unsupported file format: {ext}")
